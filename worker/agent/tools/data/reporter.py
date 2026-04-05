"""Report generation tool - create professional reports combining data, text, and charts."""

from pathlib import Path
from typing import Any

from agent.tools.base import BaseTool


class ReportGeneratorTool(BaseTool):
    name = "report_generator"
    description = (
        "Generate professional reports in PDF or Word format. "
        "Combines text sections, data tables, charts, and summaries into a polished document."
    )
    parameters = {
        "type": "object",
        "properties": {
            "format": {"type": "string", "enum": ["pdf", "docx"], "description": "Output format"},
            "output_path": {"type": "string"},
            "title": {"type": "string"},
            "author": {"type": "string"},
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "type": {"type": "string", "enum": ["heading", "paragraph", "table", "chart_image", "bullet_list"]},
                        "content": {},
                    },
                },
            },
        },
        "required": ["format", "output_path", "title", "sections"],
    }

    async def execute(self, **kwargs: Any) -> Any:
        fmt = kwargs["format"]
        output = kwargs["output_path"]
        title = kwargs["title"]
        sections = kwargs["sections"]

        Path(output).parent.mkdir(parents=True, exist_ok=True)

        if fmt == "docx":
            return self._generate_docx(output, title, kwargs.get("author", "White-Ops"), sections)
        else:
            return self._generate_pdf(output, title, kwargs.get("author", "White-Ops"), sections)

    def _generate_docx(self, output: str, title: str, author: str, sections: list) -> str:
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Author: {author}")
        doc.add_paragraph("")

        for section in sections:
            stype = section.get("type", "paragraph")
            content = section.get("content", "")

            if stype == "heading":
                level = section.get("level", 1)
                doc.add_heading(str(content), level=level)
            elif stype == "paragraph":
                doc.add_paragraph(str(content))
            elif stype == "table":
                if isinstance(content, list) and content:
                    table = doc.add_table(rows=len(content), cols=len(content[0]))
                    table.style = "Table Grid"
                    for i, row in enumerate(content):
                        for j, cell in enumerate(row):
                            table.cell(i, j).text = str(cell)
            elif stype == "chart_image":
                if Path(str(content)).exists():
                    doc.add_picture(str(content), width=Inches(6))
            elif stype == "bullet_list":
                if isinstance(content, list):
                    for item in content:
                        doc.add_paragraph(str(item), style="List Bullet")

        doc.save(output)
        return f"Report generated: {output}"

    def _generate_pdf(self, output: str, title: str, author: str, sections: list) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        from reportlab.lib.units import cm

        doc = SimpleDocTemplate(output, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(title, styles["Title"]))
        story.append(Paragraph(f"Author: {author}", styles["Normal"]))
        story.append(Spacer(1, 20))

        for section in sections:
            stype = section.get("type", "paragraph")
            content = section.get("content", "")

            if stype == "heading":
                level = section.get("level", 1)
                style = styles[f"Heading{min(level, 3)}"]
                story.append(Paragraph(str(content), style))
            elif stype == "paragraph":
                story.append(Paragraph(str(content), styles["Normal"]))
                story.append(Spacer(1, 6))
            elif stype == "table":
                if isinstance(content, list) and content:
                    t = Table(content)
                    t.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4c6ef5")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8f9fa")]),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 12))
            elif stype == "chart_image":
                if Path(str(content)).exists():
                    story.append(Image(str(content), width=16 * cm, height=10 * cm))
                    story.append(Spacer(1, 12))
            elif stype == "bullet_list":
                if isinstance(content, list):
                    for item in content:
                        story.append(Paragraph(f"\u2022 {item}", styles["Normal"]))

        doc.build(story)
        return f"Report generated: {output}"
