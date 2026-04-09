"""Word document tool - create and edit .docx files."""

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from agent.tools.base import BaseTool


class WordTool(BaseTool):
    name = "word"
    description = (
        "Create and edit Word documents (.docx). "
        "Supports: creating documents, adding paragraphs, headings, tables, "
        "images, lists, page breaks, headers/footers, and styling."
    )
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "enum": [
                    "create",
                    "read",
                    "add_heading",
                    "add_paragraph",
                    "add_table",
                    "add_image",
                    "add_list",
                    "add_page_break",
                ],
                "description": "Action to perform",
            },
            "filepath": {"type": "string", "description": "Path to the .docx file"},
            "text": {"type": "string", "description": "Text content"},
            "level": {"type": "integer", "description": "Heading level (0-9)"},
            "style": {"type": "object", "description": "Style options"},
            "data": {
                "type": "array",
                "description": "Table data (2D array)",
                "items": {"type": "array", "items": {"type": "string"}},
            },
            "image_path": {"type": "string", "description": "Path to image file"},
            "items": {
                "type": "array",
                "description": "List items",
                "items": {"type": "string"},
            },
        },
        "required": ["action", "filepath"],
    }

    async def execute(self, **kwargs: Any) -> Any:
        action = kwargs["action"]
        filepath = kwargs["filepath"]

        if action == "create":
            return self._create(filepath)
        elif action == "read":
            return self._read(filepath)
        elif action == "add_heading":
            return self._add_heading(filepath, kwargs.get("text", ""), kwargs.get("level", 1))
        elif action == "add_paragraph":
            return self._add_paragraph(filepath, kwargs.get("text", ""), kwargs.get("style"))
        elif action == "add_table":
            return self._add_table(filepath, kwargs.get("data", []))
        elif action == "add_image":
            return self._add_image(filepath, kwargs.get("image_path", ""))
        elif action == "add_list":
            return self._add_list(filepath, kwargs.get("items", []))
        elif action == "add_page_break":
            return self._add_page_break(filepath)

        return f"Unknown action: {action}"

    def _create(self, filepath: str) -> str:
        doc = Document()
        Path(filepath).parent.mkdir(parents=True, exist_ok=True)
        doc.save(filepath)
        return f"Created Word document: {filepath}"

    def _read(self, filepath: str) -> str:
        doc = Document(filepath)
        content = []
        for para in doc.paragraphs:
            content.append({"style": para.style.name, "text": para.text})
        return json.dumps({"paragraphs": len(content), "content": content})

    def _add_heading(self, filepath: str, text: str, level: int) -> str:
        doc = Document(filepath)
        doc.add_heading(text, level=level)
        doc.save(filepath)
        return f"Added heading: {text}"

    def _add_paragraph(self, filepath: str, text: str, style: dict | None) -> str:
        doc = Document(filepath)
        para = doc.add_paragraph(text)
        if style:
            if style.get("bold"):
                for run in para.runs:
                    run.bold = True
            if style.get("font_size"):
                for run in para.runs:
                    run.font.size = Pt(style["font_size"])
            if style.get("align") == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(filepath)
        return f"Added paragraph ({len(text)} chars)"

    def _add_table(self, filepath: str, data: list) -> str:
        doc = Document(filepath)
        if not data:
            return "No data provided for table"
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = "Table Grid"
        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
        doc.save(filepath)
        return f"Added table ({len(data)} rows x {len(data[0])} cols)"

    def _add_image(self, filepath: str, image_path: str) -> str:
        doc = Document(filepath)
        doc.add_picture(image_path, width=Inches(5))
        doc.save(filepath)
        return f"Added image: {image_path}"

    def _add_list(self, filepath: str, items: list) -> str:
        doc = Document(filepath)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
        doc.save(filepath)
        return f"Added list with {len(items)} items"

    def _add_page_break(self, filepath: str) -> str:
        doc = Document(filepath)
        doc.add_page_break()
        doc.save(filepath)
        return "Added page break"
